from fastapi import FastAPI, HTTPException, status
from motor.motor_asyncio import AsyncIOMotorClient
from beanie import init_beanie
from fastapi.middleware.cors import CORSMiddleware
from app.models import User, UserSubmission
from app.schemas import UserSchema
from datetime import datetime
import json
from pydantic import BaseModel
from typing import Dict, Any
from fastapi.responses import StreamingResponse
from io import BytesIO
from jinja2 import Template
import os
from docx import Document
from bson import ObjectId
from fastapi.responses import FileResponse
import os
from bs4 import BeautifulSoup

app = FastAPI()

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Define the collection for jinja_templates
jinja_templates = None

@app.on_event("startup")
async def startup():
    # Initialize the database connection
    client = AsyncIOMotorClient("mongodb://localhost:27017")
    db = client.my_database

    # Initialize collections
    global jinja_templates
    jinja_templates = db.jinja_templates  # Accessing the jinja_templates collection
    await init_beanie(database=db, document_models=[User, UserSubmission])
    
    

@app.post("/api/users")
async def create_user(user: UserSchema):
    try:
        new_user = User(
            form_fields=user.dict(),
            uiSchema=user.dict()
        )
        await new_user.insert()

        return {"message": "User created successfully!"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error creating user: {str(e)}")


@app.get("/api/users/{user_id}")
async def get_user(user_id: str):
    user = await User.get(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return {"form_fields": user.form_fields, "uiSchema": user.uiSchema}

class SubmissionData(BaseModel):
    userId: str  # User ID as a string
    submission: Dict[str, Any]  # Form submission data as a dictionary

@app.post("/api/user-submissions")
async def create_user_submission(submission_data: SubmissionData):
    user_id = submission_data.userId
    form_data = submission_data.submission

    # Ensure the user exists
    user = await User.get(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # Create a new submission
    new_submission = UserSubmission(
        user_id=user_id,
        submission=form_data,
    )
    await new_submission.insert()

    return {"message": "Submission stored successfully!"}

@app.get("/api/user-submissions/{user_id}")
async def get_user_submissions(user_id: str):
    # Ensure the user exists
    user = await User.get(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # Fetch all submissions related to the user
    submissions = await UserSubmission.find({"user_id": user_id}).to_list()

    if not submissions:
        raise HTTPException(status_code=404, detail="No submissions found for this user")

    return {"user_id": user_id, "submissions": submissions}



@app.get("/api/generate-word/{submission_id}")
async def generate_word(submission_id: str):
    try:
        # Fetch the user submission based on the submission ID
        submission = await UserSubmission.get(submission_id)
        if not submission:
            raise HTTPException(status_code=404, detail="Submission not found")

        # Fetch the template from the database
        template_document = await jinja_templates.find_one({"template_name": "user_submission_template"})
        if not template_document:
            raise HTTPException(status_code=404, detail="Template not found")

        # Extract the template content
        template_content = template_document["template_content"]

        # Use Jinja2 to render the template with submission data
        jinja_template = Template(template_content)
        rendered_html = jinja_template.render(
            user_id=submission.user_id,
            submission=submission.submission,
            created_at=submission.created_at
        )

        # Strip HTML and ensure proper line breaks
        soup = BeautifulSoup(rendered_html, "html.parser")
        plain_text = soup.get_text(separator="\n")  # Add newlines between elements

        # Generate Word document from formatted text
        doc = Document()
        for line in plain_text.split("\n"):  # Add each line as a new paragraph
            doc.add_paragraph(line.strip())

        # Save the document to a BytesIO object
        word_file = BytesIO()
        doc.save(word_file)
        word_file.seek(0)

        # Return the Word file as a streaming response
        return StreamingResponse(word_file, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=user_submission.docx"})
    
    except Exception as e:
        print(f"Error: {e}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Internal Server Error: {e}")


# Helper function to generate the Word document
def generate_word_document(submission, template_content):
    # Use Jinja2 to render the template with submission data
    jinja_template = Template(template_content)
    rendered_content = jinja_template.render(
        user_id=submission.user_id,
        submission=submission.submission,
        created_at=submission.created_at
    )

    # Generate Word document from rendered content
    doc = Document()
    doc.add_paragraph(rendered_content)

    # Save the document to a BytesIO object
    word_file = BytesIO()
    doc.save(word_file)
    word_file.seek(0)

    return word_file